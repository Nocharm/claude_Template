"""docx 파일 → 단일 JSON dict 변환 오케스트레이터."""

from functools import partial
from pathlib import Path

from docx import Document

from docx_parser.images import attach_captions, extract_images_in_paragraph
from docx_parser.ns import qn
from docx_parser.paragraphs import extract_paragraph
from docx_parser.refs import extract_bookmarks, extract_refs_in_paragraph
from docx_parser.sections import extract_sections
from docx_parser.tables import extract_table


def parse_docx(path: str | Path, *, include_raw_xml: bool = False) -> dict:
    """docx 파일을 읽어 구조 정보를 손실 없이 dict로 반환."""
    doc = Document(str(path))
    doc_part = doc.part
    body = doc.element.body

    numbering_part = _get_numbering_root(doc)
    style_lookup = _build_style_name_lookup(doc)
    rels_map = _build_image_rels_map(doc_part)

    para_extractor = partial(
        extract_paragraph,
        numbering_part=numbering_part,
        style_name_lookup=style_lookup,
        include_raw_xml=include_raw_xml,
    )

    paragraphs: list[dict] = []
    images: list[dict] = []
    tables: list[dict] = []
    refs: list[dict] = []
    bookmarks: list[dict] = []

    para_idx = 0
    for child in body:
        if child.tag == qn("w:p"):
            para = extract_paragraph(
                child,
                para_idx,
                numbering_part=numbering_part,
                style_name_lookup=style_lookup,
                include_raw_xml=include_raw_xml,
            )
            paragraphs.append(para)
            images.extend(extract_images_in_paragraph(child, para_idx, rels_map))
            refs.extend(extract_refs_in_paragraph(child, para_idx))
            bookmarks.extend(extract_bookmarks(child, para_idx))
            para_idx += 1
        elif child.tag == qn("w:tbl"):
            anchor_para_idx = para_idx - 1  # 표 직전 단락 인덱스 (음수면 문서 시작)
            tbl = extract_table(
                child,
                table_index=len(tables),
                anchor_paragraph_index=anchor_para_idx,
                para_extractor=lambda p, i: extract_paragraph(
                    p, i,
                    numbering_part=numbering_part,
                    style_name_lookup=style_lookup,
                    include_raw_xml=include_raw_xml,
                ),
            )
            # 표 내부의 그림/참조도 수집 (앵커는 표 직전 단락 인덱스로 통일)
            images.extend(extract_images_in_paragraph(child, anchor_para_idx, rels_map))
            tables.append(tbl)

    attach_captions(images, tables, paragraphs)

    sections = extract_sections(
        body,
        doc_part,
        para_extractor=lambda p, i: extract_paragraph(
            p, i,
            numbering_part=numbering_part,
            style_name_lookup=style_lookup,
            include_raw_xml=include_raw_xml,
        ),
    )

    return {
        "source_path": str(path),
        "paragraphs": paragraphs,
        "images": images,
        "tables": tables,
        "cross_references": refs,
        "bookmarks": bookmarks,
        "sections": sections,
    }


def _get_numbering_root(doc):
    try:
        return doc.part.numbering_part.element  # type: ignore[union-attr]
    except (AttributeError, KeyError):
        return None


def _build_style_name_lookup(doc) -> dict[str, str]:
    """style_id → 표시명. python-docx Styles 컬렉션 활용."""
    lookup: dict[str, str] = {}
    for style in doc.styles:
        try:
            sid = style.style_id
            name = style.name
        except (AttributeError, KeyError):
            continue
        if sid:
            lookup[sid] = name
    return lookup


def _build_image_rels_map(doc_part) -> dict[str, str]:
    """rId → target path (이미지 파일 경로)."""
    out: dict[str, str] = {}
    for rid, rel in doc_part.rels.items():
        if "image" in rel.reltype.lower():
            out[rid] = rel.target_ref
    return out
