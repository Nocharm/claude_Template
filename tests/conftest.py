"""테스트용 fixture: 합성 .docx를 즉석에서 만든다."""

import struct
import zlib
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn as docx_qn
from docx.shared import Inches, Pt
from lxml import etree


def _png_1x1() -> bytes:
    """가장 작은 유효 1x1 RGBA PNG 바이트."""
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 6, 0, 0, 0)
    ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", zlib.crc32(b"IHDR" + ihdr_data))
    raw = b"\x00\xff\x00\x00\xff"  # filter byte + RGBA
    idat_data = zlib.compress(raw)
    idat = struct.pack(">I", len(idat_data)) + b"IDAT" + idat_data + struct.pack(">I", zlib.crc32(b"IDAT" + idat_data))
    iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", zlib.crc32(b"IEND"))
    return sig + ihdr + idat + iend


@pytest.fixture
def docx_factory(tmp_path):
    """build(): docx.Document → save → 경로 반환."""
    counter = {"n": 0}

    def _save(doc) -> Path:
        counter["n"] += 1
        path = tmp_path / f"sample_{counter['n']}.docx"
        doc.save(str(path))
        return path

    return _save


@pytest.fixture
def add_soft_break():
    """단락에 soft break(Shift+Enter)를 삽입하는 헬퍼."""

    def _add(paragraph, text_before: str, text_after: str):
        p = paragraph._p
        # 기존 run 모두 제거
        for r in p.findall(docx_qn("w:r")):
            p.remove(r)
        # run1: 앞 텍스트
        r1 = etree.SubElement(p, docx_qn("w:r"))
        t1 = etml_t(r1)
        t1.text = text_before
        # soft break
        etree.SubElement(r1, docx_qn("w:br"))
        # run2: 뒤 텍스트
        r2 = etree.SubElement(p, docx_qn("w:r"))
        t2 = etml_t(r2)
        t2.text = text_after

    return _add


def etml_t(parent):
    t = etree.SubElement(parent, docx_qn("w:t"))
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return t


@pytest.fixture
def make_full_docx(docx_factory, add_soft_break):
    """검증 기준을 모두 커버하는 합성 docx 1건 생성."""

    def _make() -> Path:
        doc = Document()

        # 0: 자동 번호 단락 (List Number 스타일)
        try:
            p0 = doc.add_paragraph("자동 번호 항목 A", style="List Number")
        except KeyError:
            p0 = doc.add_paragraph("자동 번호 항목 A")
        # 1: 수동 번호 단락
        doc.add_paragraph("(1) 수동으로 입력한 번호")
        # 2: 한글 수동 번호
        doc.add_paragraph("가. 한글 글머리")
        # 3: 영문 본문
        doc.add_paragraph("This is a normal paragraph.")
        # 4: soft break 포함 단락
        p4 = doc.add_paragraph()
        add_soft_break(p4, "첫 줄", "둘째 줄")
        # 5: 캡션 누락 그림 단락 (그림 + 다음 단락이 캡션 아님)
        run = doc.add_paragraph().add_run()
        # 1x1 PNG 임베드
        from io import BytesIO
        png_bytes = _png_1x1()
        run.add_picture(BytesIO(png_bytes), width=Inches(1))
        doc.add_paragraph("그림 다음 일반 단락 (캡션 아님)")

        # 6~7: 캡션이 직후에 오는 그림
        run2 = doc.add_paragraph().add_run()
        run2.add_picture(BytesIO(_png_1x1()), width=Inches(1))
        try:
            doc.add_paragraph("그림 1 첫 번째 도식", style="Caption")
        except KeyError:
            doc.add_paragraph("그림 1 첫 번째 도식")

        # 8: 표
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "A"
        tbl.cell(0, 1).text = "B"
        tbl.cell(1, 0).text = "C"
        tbl.cell(1, 1).text = "D"

        # 섹션: 마지막 섹션을 가로(landscape)로 변경
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        return docx_factory(doc)

    return _make
