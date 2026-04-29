"""6개 합성 .docx fixture 생성기.

목적: SOP 패턴(한/영 혼용, 다단계 목록, 캡션, 다중 섹션, 참조)을 재현하여
파서 검증 + 문서 예시 + 후속 로직 테스트용 입력으로 활용.

사용:
    python tests/fixtures/generate_inputs.py
"""

from __future__ import annotations

import struct
import zlib
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree


XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"
OUT_DIR = Path(__file__).parent / "inputs"


# --- 공통 헬퍼 -------------------------------------------------------------

def _png_1x1() -> bytes:
    sig = b"\x89PNG\r\n\x1a\n"
    ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 6, 0, 0, 0)
    ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", zlib.crc32(b"IHDR" + ihdr_data))
    raw = b"\x00\xff\x00\x00\xff"
    idat_data = zlib.compress(raw)
    idat = struct.pack(">I", len(idat_data)) + b"IDAT" + idat_data + struct.pack(">I", zlib.crc32(b"IDAT" + idat_data))
    iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", zlib.crc32(b"IEND"))
    return sig + ihdr + idat + iend


def add_multilevel_numbering(doc, abstract_id: int = 99, num_id: int = 99) -> int:
    """numbering.xml에 3단계 번호 정의 등록 후 numId 반환.

    레벨 0: 1.   2.   3.       (decimal)
    레벨 1: 1)   2)   3)       (decimalParen, lvlText "%2)")
    레벨 2: (a)  (b)  (c)      (lowerLetter)
    """
    numbering = doc.part.numbering_part.element

    abstract = etree.SubElement(numbering, qn("w:abstractNum"))
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    levels = [
        ("decimal", "%1.", 0),
        ("decimal", "%2)", 720),
        ("lowerLetter", "(%3)", 1440),
    ]
    for ilvl, (fmt, txt, indent) in enumerate(levels):
        lvl = etree.SubElement(abstract, qn("w:lvl"))
        lvl.set(qn("w:ilvl"), str(ilvl))
        start = etree.SubElement(lvl, qn("w:start"))
        start.set(qn("w:val"), "1")
        numFmt = etree.SubElement(lvl, qn("w:numFmt"))
        numFmt.set(qn("w:val"), fmt)
        lvlText = etree.SubElement(lvl, qn("w:lvlText"))
        lvlText.set(qn("w:val"), txt)
        pPr = etree.SubElement(lvl, qn("w:pPr"))
        ind = etree.SubElement(pPr, qn("w:ind"))
        ind.set(qn("w:left"), str(indent))

    num = etree.SubElement(numbering, qn("w:num"))
    num.set(qn("w:numId"), str(num_id))
    abs_ref = etree.SubElement(num, qn("w:abstractNumId"))
    abs_ref.set(qn("w:val"), str(abstract_id))

    return num_id


def apply_numbering(p, num_id: int, ilvl: int = 0) -> None:
    pPr = p._p.get_or_add_pPr()
    numPr = etree.SubElement(pPr, qn("w:numPr"))
    ilvl_el = etree.SubElement(numPr, qn("w:ilvl"))
    ilvl_el.set(qn("w:val"), str(ilvl))
    numId_el = etree.SubElement(numPr, qn("w:numId"))
    numId_el.set(qn("w:val"), str(num_id))


def add_soft_break_paragraph(doc, *segments: str):
    """세그먼트들을 Shift+Enter(soft break)로 연결한 단락 추가."""
    p = doc.add_paragraph()
    for i, seg in enumerate(segments):
        r = etree.SubElement(p._p, qn("w:r"))
        if i > 0:
            etree.SubElement(r, qn("w:br"))
        t = etree.SubElement(r, qn("w:t"))
        t.set(XML_SPACE, "preserve")
        t.text = seg
    return p


def add_picture(doc):
    """1.5인치 폭의 1x1 PNG 그림을 새 단락에 인라인 삽입."""
    p = doc.add_paragraph()
    p.add_run().add_picture(BytesIO(_png_1x1()), width=Inches(1.5))
    return p


def add_bookmark(p, bm_id: int, name: str) -> None:
    bm_start = etree.Element(qn("w:bookmarkStart"))
    bm_start.set(qn("w:id"), str(bm_id))
    bm_start.set(qn("w:name"), name)
    p._p.insert(0, bm_start)
    bm_end = etree.SubElement(p._p, qn("w:bookmarkEnd"))
    bm_end.set(qn("w:id"), str(bm_id))


def add_ref_field(p, target_bookmark: str, display_text: str) -> None:
    """w:fldSimple 형식 REF 필드 (표시 텍스트는 캐시된 값)."""
    fld = etree.SubElement(p._p, qn("w:fldSimple"))
    fld.set(qn("w:instr"), f" REF {target_bookmark} \\h ")
    r = etree.SubElement(fld, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.text = display_text


def caption_style_or_none(doc) -> str | None:
    """Caption 스타일이 있으면 그 이름, 없으면 None (스타일 미적용)."""
    for s in doc.styles:
        if s.name == "Caption":
            return "Caption"
    return None


# --- 케이스 빌더 -----------------------------------------------------------

def build_normal() -> Path:
    """정상: 다단계 자동 번호 + 캡션 정상 부여 그림/표, 단일 portrait 섹션."""
    doc = Document()
    doc.add_paragraph("표준운영절차 SOP-001", style="Heading 1")
    doc.add_paragraph("This document describes the standard preparation procedure.")

    num_id = add_multilevel_numbering(doc)

    p = doc.add_paragraph("작업 준비 (Preparation)")
    apply_numbering(p, num_id, ilvl=0)
    p = doc.add_paragraph("Verify equipment status")
    apply_numbering(p, num_id, ilvl=1)
    p = doc.add_paragraph("Check power supply")
    apply_numbering(p, num_id, ilvl=2)
    p = doc.add_paragraph("작업 수행 (Execution)")
    apply_numbering(p, num_id, ilvl=0)

    add_picture(doc)
    cap_style = caption_style_or_none(doc)
    if cap_style:
        doc.add_paragraph("그림 1: System overview", style=cap_style)
    else:
        doc.add_paragraph("그림 1: System overview")

    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Item"
    tbl.cell(0, 1).text = "Status"
    tbl.cell(1, 0).text = "전원"
    tbl.cell(1, 1).text = "OK"
    if cap_style:
        doc.add_paragraph("표 1: Equipment checklist", style=cap_style)
    else:
        doc.add_paragraph("표 1: Equipment checklist")

    out = OUT_DIR / "normal.docx"
    doc.save(str(out))
    return out


def build_mixed_numbering() -> Path:
    """자동 + 수동 번호 혼재. 동일 위계인데 일부만 자동인 케이스 포함."""
    doc = Document()
    doc.add_paragraph("작업 절차", style="Heading 1")

    num_id = add_multilevel_numbering(doc)

    p = doc.add_paragraph("준비 단계")
    apply_numbering(p, num_id, ilvl=0)  # 자동 1.
    p = doc.add_paragraph("실행 단계")
    apply_numbering(p, num_id, ilvl=0)  # 자동 2.

    # 수동 번호가 자동 흐름 사이에 끼어든 케이스
    doc.add_paragraph("3. 점검 단계 (수동 입력)")          # 수동 데시멀
    doc.add_paragraph("① 첫 번째 점검 항목")               # 수동 원문자
    doc.add_paragraph("② 두 번째 점검 항목")
    doc.add_paragraph("가. 한글 글머리")
    doc.add_paragraph("(1) 괄호 번호")
    doc.add_paragraph("- 대시 글머리 항목")

    p = doc.add_paragraph("종료 단계")
    apply_numbering(p, num_id, ilvl=0)  # 자동 — but 번호 재시작 가능성 있음

    out = OUT_DIR / "mixed_numbering.docx"
    doc.save(str(out))
    return out


def build_soft_break() -> Path:
    """soft return vs hard return 비교용."""
    doc = Document()
    doc.add_paragraph("줄바꿈 비교 케이스", style="Heading 1")

    # soft break 단락 (한 단락 내 2줄)
    add_soft_break_paragraph(doc, "첫 번째 줄 (soft break 위)", "두 번째 줄 (soft break 아래)")

    # hard return 으로 같은 내용 (별도 단락 2개)
    doc.add_paragraph("첫 번째 줄 (hard return 위)")
    doc.add_paragraph("두 번째 줄 (hard return 아래)")

    # soft break 3개 들어간 단락
    add_soft_break_paragraph(
        doc,
        "Step 1: Initialize the system",
        "Step 2: Validate inputs",
        "Step 3: Execute operation",
        "Step 4: Cleanup resources",
    )

    out = OUT_DIR / "soft_break.docx"
    doc.save(str(out))
    return out


def build_missing_caption() -> Path:
    """캡션 누락 케이스. 명시적 null 처리 검증용."""
    doc = Document()
    doc.add_paragraph("캡션 누락 검증", style="Heading 1")

    # 1. 캡션 없음 (그림 직후 일반 단락)
    add_picture(doc)
    doc.add_paragraph("이것은 그림에 대한 설명이 아니라 다음 본문이다.")

    # 2. 캡션 있음 (대조용)
    add_picture(doc)
    cap_style = caption_style_or_none(doc)
    if cap_style:
        doc.add_paragraph("그림 1: 정상 캡션", style=cap_style)
    else:
        doc.add_paragraph("그림 1: 정상 캡션")

    # 3. 캡션 없음 (그림 다음에 또 다른 그림)
    add_picture(doc)
    add_picture(doc)
    doc.add_paragraph("두 그림 사이에 캡션이 없는 케이스.")

    # 4. 표 캡션 누락
    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "K"
    tbl.cell(0, 1).text = "V"
    doc.add_paragraph("표 직후 일반 본문 (캡션 아님).")

    out = OUT_DIR / "missing_caption.docx"
    doc.save(str(out))
    return out


def build_multi_section() -> Path:
    """portrait → landscape → portrait, 섹션별 머리글/꼬리글 + 첫 페이지 머리글."""
    doc = Document()

    # 섹션 1 (portrait)
    s1 = doc.sections[0]
    s1.different_first_page_header_footer = True
    s1.first_page_header.add_paragraph("섹션 1 첫 페이지 헤더")
    s1.header.add_paragraph("섹션 1 일반 헤더")
    s1.footer.add_paragraph("섹션 1 푸터")
    doc.add_paragraph("섹션 1 본문 (portrait)", style="Heading 1")
    doc.add_paragraph("This is the first portrait section.")

    # 섹션 2 (landscape)
    s2 = doc.add_section(WD_SECTION.NEW_PAGE)
    s2.orientation = WD_ORIENT.LANDSCAPE
    s2.page_width, s2.page_height = s1.page_height, s1.page_width
    s2.header.is_linked_to_previous = False
    s2.footer.is_linked_to_previous = False
    s2.header.add_paragraph("섹션 2 헤더 (landscape)")
    s2.footer.add_paragraph("섹션 2 푸터")
    doc.add_paragraph("섹션 2 본문 (landscape)", style="Heading 1")
    doc.add_paragraph("이 섹션은 가로 방향이다.")

    # 섹션 3 (portrait 복귀)
    s3 = doc.add_section(WD_SECTION.NEW_PAGE)
    s3.orientation = WD_ORIENT.PORTRAIT
    s3.page_width, s3.page_height = s1.page_width, s1.page_height
    s3.header.is_linked_to_previous = False
    s3.footer.is_linked_to_previous = False
    s3.header.add_paragraph("섹션 3 헤더")
    s3.footer.add_paragraph("섹션 3 푸터")
    doc.add_paragraph("섹션 3 본문 (portrait)", style="Heading 1")
    doc.add_paragraph("portrait 복귀 후 마지막 섹션.")

    out = OUT_DIR / "multi_section.docx"
    doc.save(str(out))
    return out


def build_cross_reference() -> Path:
    """그림/표/제목 참조(REF 필드) + 북마크."""
    doc = Document()
    doc.add_paragraph("참조 검증 문서", style="Heading 1")

    # 북마크 달린 제목
    h2 = doc.add_paragraph("핵심 절차", style="Heading 2")
    add_bookmark(h2, bm_id=1, name="_Ref_core_section")

    # 그림 + 캡션 (캡션 단락에 북마크)
    add_picture(doc)
    cap_style = caption_style_or_none(doc)
    fig_cap = doc.add_paragraph()
    if cap_style:
        fig_cap.style = cap_style
    fig_cap.add_run("그림 1: 핵심 도식")
    add_bookmark(fig_cap, bm_id=2, name="_Ref_fig_1")

    # 표 + 캡션 (캡션 단락에 북마크)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.cell(0, 0).text = "data"
    tbl_cap = doc.add_paragraph()
    if cap_style:
        tbl_cap.style = cap_style
    tbl_cap.add_run("표 1: 데이터 요약")
    add_bookmark(tbl_cap, bm_id=3, name="_Ref_tbl_1")

    # 본문에 REF 삽입
    body = doc.add_paragraph("자세한 내용은 ")
    add_ref_field(body, "_Ref_core_section", "핵심 절차")
    body.add_run(" 절을 참고하라. ")
    add_ref_field(body, "_Ref_fig_1", "그림 1")
    body.add_run(" 및 ")
    add_ref_field(body, "_Ref_tbl_1", "표 1")
    body.add_run(" 참조.")

    out = OUT_DIR / "cross_reference.docx"
    doc.save(str(out))
    return out


# --- 엔트리 ----------------------------------------------------------------

BUILDERS = [
    build_normal,
    build_mixed_numbering,
    build_soft_break,
    build_missing_caption,
    build_multi_section,
    build_cross_reference,
]


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for build in BUILDERS:
        path = build()
        print(f"  generated: {path.relative_to(OUT_DIR.parent.parent)}")


if __name__ == "__main__":
    main()
